#!/usr/bin/env python3
"""Scan anonymized submission files for forbidden identifying terms.

Usage:
    python scripts/check_blind_submission.py
    python scripts/check_blind_submission.py --path submission/springer_ailaw/anonymous_artifact/
    python scripts/check_blind_submission.py --path submission/springer_ailaw/anonymized_manuscript/
    python scripts/check_blind_submission.py --all   # include public files
"""

import argparse
import re
import sys
from pathlib import Path

REPO_ROOT = Path(__file__).resolve().parent.parent

DEFAULT_PATHS = [
    REPO_ROOT / "submission" / "springer_ailaw" / "anonymized_manuscript",
    REPO_ROOT / "submission" / "springer_ailaw" / "anonymous_artifact",
    REPO_ROOT / "docs" / "data_availability_statement_anonymous.md",
]

TEXT_EXTENSIONS = {
    ".md", ".txt", ".py", ".json", ".jsonl", ".csv",
    ".yaml", ".yml", ".toml", ".ttl", ".rego", ".rst",
}

# Each entry: (pattern_str, use_word_boundary, severity)
FORBIDDEN_TERMS = [
    ("Roy Saurabh",         False, "FAIL"),
    ("Saurabh",             False, "FAIL"),
    ("AffectLog SAS",       False, "FAIL"),
    ("AffectLog",           False, "FAIL"),
    ("roy@affectlog.com",   False, "FAIL"),
    ("github.com/roy-saurabh", False, "FAIL"),
    ("zenodo.20399201",     False, "FAIL"),
    ("20399201",            False, "FAIL"),
    ("20397383",            False, "FAIL"),
    ("ORCID",               False, "FAIL"),
    ("0000-0003-3439-7731", False, "FAIL"),
    ("founder of AffectLog",False, "FAIL"),
    # "Roy" as whole word — warn but don't hard-fail (could be legitimate English)
    ("Roy",                 True,  "WARN"),
]


def _build_regex(term: str, word_boundary: bool) -> re.Pattern:
    escaped = re.escape(term)
    if word_boundary:
        return re.compile(r"\b" + escaped + r"\b", re.IGNORECASE)
    return re.compile(escaped, re.IGNORECASE)


COMPILED = [(t, _build_regex(t, wb), sev) for t, wb, sev in FORBIDDEN_TERMS]


def _check_text_lines(lines: list[str], path_label: str) -> list[dict]:
    findings = []
    for lineno, line in enumerate(lines, 1):
        for term, pattern, severity in COMPILED:
            if pattern.search(line):
                findings.append({
                    "path": path_label,
                    "term": term,
                    "location": f"line {lineno}",
                    "severity": severity,
                    "snippet": line.strip()[:120],
                })
    return findings


def _check_docx(fpath: Path, path_label: str) -> list[dict]:
    try:
        import docx  # type: ignore
    except ImportError:
        return []
    findings = []
    doc = docx.Document(str(fpath))
    items = [("para", i + 1, p.text) for i, p in enumerate(doc.paragraphs)]
    for tbl_idx, tbl in enumerate(doc.tables):
        for row_idx, row in enumerate(tbl.rows):
            for cell_idx, cell in enumerate(row.cells):
                items.append((f"table{tbl_idx+1}/row{row_idx+1}/cell{cell_idx+1}", 0, cell.text))
    for loc_type, loc_num, text in items:
        for term, pattern, severity in COMPILED:
            if pattern.search(text):
                location = f"{loc_type} {loc_num}" if loc_num else loc_type
                findings.append({
                    "path": path_label,
                    "term": term,
                    "location": location,
                    "severity": severity,
                    "snippet": text.strip()[:120],
                })
    return findings


def scan_path(target: Path) -> list[dict]:
    findings = []
    if not target.exists():
        print(f"[skip] path not found: {target}", file=sys.stderr)
        return findings
    files = [target] if target.is_file() else sorted(target.rglob("*"))
    for fpath in files:
        if not fpath.is_file():
            continue
        label = str(fpath.relative_to(REPO_ROOT))
        ext = fpath.suffix.lower()
        if ext == ".docx":
            findings.extend(_check_docx(fpath, label))
        elif ext in TEXT_EXTENSIONS:
            try:
                lines = fpath.read_text(encoding="utf-8", errors="replace").splitlines()
            except Exception:
                continue
            findings.extend(_check_text_lines(lines, label))
    return findings


def print_table(findings: list[dict]) -> None:
    if not findings:
        print("  (no findings)")
        return
    col_path  = max(len(f["path"])     for f in findings)
    col_term  = max(len(f["term"])     for f in findings)
    col_loc   = max(len(f["location"]) for f in findings)
    col_sev   = 4
    header = (
        f"{'PATH':<{col_path}}  {'TERM':<{col_term}}  "
        f"{'LOCATION':<{col_loc}}  {'SEV':<{col_sev}}"
    )
    print(header)
    print("-" * len(header))
    for f in findings:
        print(
            f"{f['path']:<{col_path}}  {f['term']:<{col_term}}  "
            f"{f['location']:<{col_loc}}  {f['severity']:<{col_sev}}"
        )
        print(f"  snippet: {f['snippet']}")


def main() -> None:
    parser = argparse.ArgumentParser(description="Blind-submission checker")
    parser.add_argument(
        "--path", type=Path, default=None,
        help="Specific path to scan (file or directory)",
    )
    parser.add_argument(
        "--all", action="store_true",
        help="Also scan public (non-anonymous) files",
    )
    args = parser.parse_args()

    if args.path:
        paths = [args.path]
    elif args.all:
        paths = [REPO_ROOT]
    else:
        paths = DEFAULT_PATHS

    all_findings: list[dict] = []
    for p in paths:
        found = scan_path(p)
        all_findings.extend(found)

    fails = [f for f in all_findings if f["severity"] == "FAIL"]
    warns = [f for f in all_findings if f["severity"] == "WARN"]

    print(f"\nBlind-submission check: {len(fails)} FAIL(s), {len(warns)} WARN(s)\n")
    print_table(all_findings)

    if fails:
        print(f"\nFAILED — {len(fails)} forbidden identifier(s) found.")
        sys.exit(1)
    elif warns:
        print(f"\nPASSED (with {len(warns)} warning(s)).")
    else:
        print("\nPASSED — no forbidden identifiers found.")


if __name__ == "__main__":
    main()
